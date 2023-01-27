import pandas as pd
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from docx.shared import RGBColor
import re

import romanize3
# abrindo o arquivo excel

df = pd.read_excel('Kefa I_output.xlsx')

print(df)
# criando o documento word
document = Document()
def formata(str, paragrafo,rgb):
    # Remove all tags except "i", "br" and "n"

    text = str

    text = re.sub(r'<a[^>]*>([^<]+)</a>', '\\1', text)
    text = re.sub(r'<sup[^>]*>([^<]+)</sup>', '\\1', text)


   # text = "S HANGED IS A קללת אלהים — i.e., a degradation of the Divine King"
    #r = romanize3.__dict__['heb']
    #transliterated_text = r.convert(text)
   # print(transliterated_text)

    print("0i0i0i0i0i "+text)

   # texto = "ola <br/> <n> estou aqui </n> vc nao sabe"
    segmentado = re.split(r'(<[^>]*>)', text)
    print(segmentado)



    paragrafo.paragraph_format.left_indent = Inches(0.5)
    tagi = False
    tagbr = False
    tagn = False
    print(segmentado)
    for i in segmentado:
        print(i)
       # print(rashi_paragraph)
        if "<" in i:
            run = paragrafo.add_run("")
        else:
            run = paragrafo.add_run(i)
        run.font.color.rgb = rgb
        if "<i" in i:
            tagi = True

        if "<br/>"  in i:

            run.add_break()
        if "<br>" in i:
            run.add_break()

        if "<b>" in i:
            tagn = True


        if "</i>" in i:
            tagi = False

        if "</b>" in i:
            tagn = False

        if tagi:
            print('1111')
            run.italic = True
        else:
            print('2222')
            run.italic = False
        if tagn:
            print('3333')
            run.bold = True
        else:
            print('4444')
            run.bold = False

    return paragrafo;
        #run.font.color.rgb = RGBColor(0, 0, 255)
# iterando pelas linhas do dataframe

for index, row in df.iterrows():
    # adicionando a primeira linha com o livro, capítulo e versículo
    paragraph = document.add_paragraph(f'{row["Livro"]} {row["Capitulo"]}:{row["Versiculo"]}')
    paragraph.runs[0].font.size = 14
    paragraph.runs[0].bold=True
    # adicionando a segunda linha com o conteúdo da coluna nvi
    document.add_paragraph(row["nvi"])
    # adicionando a terceira linha com o conteúdo da coluna texto1 em itálico
    texto1_paragraph = document.add_paragraph(row["texto1"])
    texto1_paragraph.italic = True

    if row["Rashi"] != None and not pd.isna(row["Rashi"])  :

        paragraph = document.add_paragraph()
        paragraph= formata(row["Rashi"], paragraph,RGBColor(59,90,142))
       # paragraph.runs[0].font.color.rgb = RGBColor(59,90,142 )
        #rashi_paragraph.paragraph_format.left_indent = Inches(0.5)
        #rashi_paragraph.italic = True
        #rashi_paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 255)
        paragraph.add_run("(Rashi)")

    if row["Ramban"] != None and not pd.isna(row["Ramban"]) :
        paragraph = document.add_paragraph()
        paragraph = formata(row["Ramban"], paragraph,RGBColor(17, 90, 107))

        paragraph.add_run("(Ramban)")

    if row["Sforno"] != None and not pd.isna(row["Sforno"]) :
        paragraph = document.add_paragraph()
        paragraph = formata(row["Sforno"], paragraph,RGBColor(107, 94, 155))
       # paragraph.runs[0].font.color.rgb = RGBColor(107, 94, 155)
        paragraph.add_run("(Sforno)")

    if row["Ibn"] != None and not pd.isna(row["Ibn"]):
        paragraph = document.add_paragraph()
        paragraph = formata(row["Ibn"], paragraph,RGBColor(30, 106, 57))
       # paragraph.runs[0].font.color.rgb = RGBColor(30, 106, 57)
        paragraph.add_run("(Ibn Ezra)")
    if row["Onkelos"] != None and not pd.isna(row["Onkelos"]):
        paragraph = document.add_paragraph()
        paragraph = formata(row["Onkelos"], paragraph,RGBColor(30, 106, 57))
       # paragraph.runs[0].font.color.rgb = RGBColor(30, 106, 57)
        paragraph.add_run("(Onkelos)")
    if row["Jonathan"] != None and not pd.isna(row["Jonathan"]):
        paragraph = document.add_paragraph()
        paragraph = formata(row["Jonathan"], paragraph,RGBColor(30, 106, 57))
       # paragraph.runs[0].font.color.rgb = RGBColor(30, 106, 57)
        paragraph.add_run("(Jonathan)")



# salvando o documento
document.save('Kefa I_output.docx')
